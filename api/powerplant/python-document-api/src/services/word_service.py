from docx import Document
import os

def create_word_document(content, filename):
    doc = Document()
    doc.add_heading('Document Title', level=1)
    doc.add_paragraph(content)

    output_dir = os.path.join(os.path.dirname(__file__), '../../output/word_files')
    os.makedirs(output_dir, exist_ok=True)

    file_path = os.path.join(output_dir, filename)
    doc.save(file_path)
    return file_path