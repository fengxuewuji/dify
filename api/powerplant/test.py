from flask import Flask, request, jsonify
from sqlalchemy import create_engine, text
import pymysql
from sqlalchemy.orm import sessionmaker
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy import Column, Integer, Float
import requests
import json
import time
import pandas as pd
import numpy as np
import yaml
import re
import difflib
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.units import cm
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import io
from flask import send_file

HCFX_TAG_NAMES = {
        'JCYY:U4_MH_GLXL': '锅炉效率耗差(g/kWh)',
        'JCYY:U4_MH_JCRH': '汽轮机基础热耗率耗差(g/kWh)',
        'JCYY:U4_MH_GTF': '高调阀运行方式耗差(g/kWh)',
        'JCYY:U4_MH_ZZQWD': '主汽温度耗差(g/kWh)',
        'JCYY:U4_MH_ZRZQWD': '再热汽温度耗差(g/kWh)',
        'JCYY:U4_MH_GRJWS': '过热减温水流量耗差(g/kWh)',
        'JCYY:U4_MH_ZRJWS': '再热减温水流量耗差(g/kWh)',
        'JCYY:U4_MH_YXBY': '运行背压耗差(g/kWh)',
        'JCYY:U4_MH_1GJSDC': '1号高加上端差耗差(g/kWh)',
        'JCYY:U4_MH_1GJXDC': '1号高加下端差耗差(g/kWh)',
        'JCYY:U4_MH_2GJSDC': '2号高加上端差耗差(g/kWh)',
        'JCYY:U4_MH_2GJXDC': '2号高加下端差耗差(g/kWh)',
        'JCYY:U4_MH_3GJSDC': '3号高加上端差耗差(g/kWh)',
        'JCYY:U4_MH_3GJXDC': '3号高加下端差耗差(g/kWh)',
        'JCYY:U4_MH_5DJSDC': '5号低加上端差耗差(g/kWh)',
        'JCYY:U4_MH_5DJXDC': '5号低加下端差耗差(g/kWh)',
        'JCYY:U4_MH_6DJSDC': '6号低加上端差耗差(g/kWh)',
        'JCYY:U4_MH_6DJXDC': '6号低加下端差耗差(g/kWh)',
        'JCYY:U4_MH_7DJSDC': '7号低加上端差耗差(g/kWh)',
        'JCYY:U4_MH_7DJXDC': '7号低加下端差耗差(g/kWh)',
        'inner_leakage': '热力系统内漏耗差(g/kWh)',
        'JCYY:U4_MH_FDBSL': '发电补水率耗差(g/kWh)',
        'JCYY:U4_MH_GYGQGRB': '工业供汽供热比耗差(g/kWh)',
        'JCYY:U4_MH_CNGRB': '采暖供热比耗差(g/kWh)',
        'JCYY:U4_MH_NFQTY': '暖风器投运耗差(g/kWh)',
        'JCYY:U4_MH_LTGJLSS': '连通管蝶阀节流损失耗差(g/kWh)',
        'JCYY:U4_MH_CLXS': '出力系数耗差(g/kWh)',
        'JCYY:U4_MH_FGC': '峰谷差耗差(g/kWh)',
    }

HCFX_TAG_NAMES_REVERSE = {v: k for k, v in HCFX_TAG_NAMES.items()}

HCFX_FORMULA = { # TODO: 这里的公式需要根据实际情况进行调整,写个yaml文件？
    '锅炉效率耗差(g/kWh)': '锅炉计算效率 - 锅炉设计效率',
}
# 创建 Flask 应用实例
app = Flask(__name__)
# 配置数据库连接
DATABASE_URI = 'mysql+pymysql://root:MS_zyp1121@localhost:3306/dify_db'
try:
    engine = create_engine(DATABASE_URI)
except pymysql.Error as e:
    app.logger.error(f"Error connecting to the database: {e}")

# 定义HCFX表ORM模型
Base = declarative_base()
class HCFX(Base):
    __tablename__ = 'HCFX'
    id = Column(Integer, primary_key=True)
    boiler_eff = Column(Float)
    turbine_base_heat_rate = Column(Float)
    hp_valve = Column(Float)
    sh_temp = Column(Float)
    rh_temp = Column(Float)
    sh_de_water = Column(Float)
    rh_de_water = Column(Float)
    back_pres_ = Column(Float)
    hp_heater_1_up = Column(Float)
    hp_heater_1_down = Column(Float)
    hp_heater_2_up = Column(Float)
    hp_heater_2_down = Column(Float)
    hp_heater_3_up = Column(Float)
    hp_heater_3_down = Column(Float)
    lp_heater_5_up = Column(Float)
    lp_heater_5_down = Column(Float)
    lp_heater_6_up = Column(Float)
    lp_heater_6_down = Column(Float)
    lp_heater_7_up = Column(Float)
    lp_heater_7_down = Column(Float)
    inner_leakage = Column(Float)
    makeup_water = Column(Float)
    industrial_heat_supply = Column(Float)
    life_heat_supply = Column(Float)
    air_heater = Column(Float)
    butterfly_valve = Column(Float)
    output_coef = Column(Float)
    peak_vally_diff = Column(Float)

# 使用 @app.route 装饰器定义一个路由，该路由的路径为 /query，只接受 POST 请求
@app.route('/query', methods=['POST'])
def query_database():
    # 从请求的 JSON 数据中获取名为 sql 的字段，即客户端发送的 SQL 查询语句
    print("------------")
    sql_query = request.json.get('sql')
    print("------dddd------")
    print(sql_query)
    if not sql_query:
        return jsonify({"error": "SQL query is required"}), 400
    try:
        # 创建一个数据库连接，并使用 with 语句管理连接的生命周期
        with engine.connect() as connection:
            # 使用 connection 对象执行 SQL 查询语句，并将结果存储在 result 变量中
            result = connection.execute(text(sql_query))
            # 从 result 中获取所有查询结果，并存储在 rows 变量中
            rows = result.fetchall()
            # 获取查询结果的列名，并存储在 columns 变量中
            columns = result.keys()
            # 将查询结果转换为字典列表
            result_dict = [dict(zip(columns, row)) for row in rows]
            # 将 result_dict 转换为 JSON 格式
            return jsonify(result_dict)

    except Exception as e:
        app.logger.error(f"Error executing SQL query: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/hcfx/mock', methods=['GET'])
def get_latest_hcfx():
    Session = sessionmaker(bind=engine)
    session = Session()
    try:
        # 列名对应表
        col_names = {
            'boiler_eff': '锅炉效率耗差(g/kWh)',
            'turbine_base_heat_rate': '汽轮机基础热耗率耗差(g/kWh)',
            'hp_valve': '高调阀运行方式耗差(g/kWh)',
            'sh_temp': '主汽温度耗差(g/kWh)',
            'rh_temp': '再热汽温度耗差(g/kWh)',
            'sh_de_water': '过热减温水流量耗差(g/kWh)',
            'rh_de_water': '再热减温水流量耗差(g/kWh)',
            'back_pres_': '运行背压耗差(g/kWh)',
            'hp_heater_1_up': '1号高加上端差耗差(g/kWh)',
            'hp_heater_1_down': '1号高加下端差耗差(g/kWh)',
            'hp_heater_2_up': '2号高加上端差耗差(g/kWh)',
            'hp_heater_2_down': '2号高加下端差耗差(g/kWh)',
            'hp_heater_3_up': '3号高加上端差耗差(g/kWh)',
            'hp_heater_3_down': '3号高加下端差耗差(g/kWh)',
            'lp_heater_5_up': '5号低加上端差耗差(g/kWh)',
            'lp_heater_5_down': '5号低加下端差耗差(g/kWh)',
            'lp_heater_6_up': '6号低加上端差耗差(g/kWh)',
            'lp_heater_6_down': '6号低加下端差耗差(g/kWh)',
            'lp_heater_7_up': '7号低加上端差耗差(g/kWh)',
            'lp_heater_7_down': '7号低加下端差耗差(g/kWh)',
            'inner_leakage': '热力系统内漏耗差(g/kWh)',
            'makeup_water': '发电补水率耗差(g/kWh)',
            'industrial_heat_supply': '工业供汽供热比耗差(g/kWh)',
            'life_heat_supply': '采暖供热比耗差(g/kWh)',
            'air_heater': '暖风器投运耗差(g/kWh)',
            'butterfly_valve': '连通管蝶阀节流损失耗差(g/kWh)',
            'output_coef': '出力系数耗差(g/kWh)',
            'peak_vally_diff': '峰谷差耗差(g/kWh)',
        }
        latest_record = session.query(HCFX).order_by(HCFX.id.desc()).first()
        if latest_record:
            result = {col.name: getattr(latest_record, col.name) for col in HCFX.__table__.columns}
            result = {col_names.get(k, k): v for k, v in result.items() if k != 'id'}  # 排除 id 字段
            # 将结果转换为 JSON 格式并返回
            return jsonify(result)
        else:
            return jsonify({'error': 'No record found'}), 404
    except Exception as e:
        app.logger.error(f"Error querying HCFX: {e}")
        return jsonify({'error': str(e)}), 500
    finally:
        session.close()


@app.route('/hcfx/real', methods=['GET'])
def get_real_hcfx_data():
    tag_names = HCFX_TAG_NAMES
    # 这里的 tag_names 是一个字典，键是标签名，值是对应的中文名称
    # 调用 get_real_value 函数获取实时数据
    real_data = get_real_value(tag_names)
    return real_data

@app.route('/loss/real', methods=['GET'])
def get_real_loss_data():
    tag_names = {
            'JCYY:U4_BOILER_Q2': 'q2',
            'JCYY:U4_BOILER_Q2': 'q2',
            'JCYY:U4_BOILER_Q3': 'q3',
            'JCYY:U4_BOILER_Q4': 'q4',
            'JCYY:U4_BOILER_Q5': 'q5',
            'JCYY:U4_BOILER_Q_OTHER': 'q6',
        }
    real_data = get_real_value(tag_names)
    return real_data

@app.route('/xnjs/real', methods=['GET'])
def get_real_xnjs_data():
    tag_names = {
            'JCYY:U4_BOILER_EFF': '锅炉计算效率(%)',
            'JCYY:U4_GYGXL': '汽轮机高压缸效率(%)',
        }
    real_data = get_real_value(tag_names)
    return real_data

def parse_json_response(response_text):
    """
    Parses a JSON response text and returns a dictionary.
    
    Parameters:
    response_text (str): The JSON response text to parse.
    
    Returns:
    dict: Parsed dictionary from the JSON response text.
    """
    try:
        return json.loads(response_text)
    except json.JSONDecodeError as e:
        print(f"Error parsing JSON response: {e}")
        return None

def get_token():
    """
    Retrieves a token from the specified URL.
    
    Returns:
    str: The token if the request is successful, otherwise None.
    """
    url = 'http://168.168.10.82:20001/v1/login'

    headers = {
        'Content-Type': 'application/json',
    }

    data = {
        "dbName":"db102",
        "userName":"user1",
        "password":"luculent1!"
    }

    response = requests.post(url, json=data, headers=headers)
    token = response.json()['token']
    return token

def get_real_value(tag_names):

    url = 'http://168.168.10.82:20001/v1/getRealtimeValue'

    token = get_token()

    headers = {
        'Content-Type': 'application/json',
        'token': token,
    }

    data = {
        "dbName":"db102",
        "tagNames": list(tag_names.keys()),
    }

    response = requests.post(url, json=data, headers=headers)
    parsed_response = parse_json_response(response.text)

    all_data = {}
    if not parsed_response or 'data' not in parsed_response:
        return {}
        
    for i in parsed_response['data']:
        all_data.update(i)

    real_data = {}
    for k, v in tag_names.items():
        if k in all_data:
            real_data[v] = all_data[k]
        else:
            real_data[v] = None  # 如果没有找到对应的标签，设置为 None
    return real_data

@app.route('/common/real', methods=['POST'])
def get_history_value():

    tag_name = request.json.get('tagName')
    end_time_str = request.json.get('endTimeStr')
    end_time = int(time.mktime(time.strptime(end_time_str, '%Y-%m-%d %H:%M:%S'))) * 1000  # 转换为毫秒
    time_span = int(request.json.get('timeSpan'))
    start_time = end_time - time_span
    # 指定返回数据的数量而非时间跨度,默认返回12条数据。
    data_num = int(request.json.get('dataNum', 12))

    if data_num < 1 or data_num > 100:
        return jsonify({"error": "dataNum must be between 1 and 100"}), 400

    resamplePriodMs = int(end_time - start_time) // data_num
    resamplePriodMs = max(resamplePriodMs, 1000)  # 确保周期至少为1秒
    
    url = 'http://168.168.10.82:20001/v1/getHisResampleValue'

    token = get_token()

    headers = {
        'Content-Type': 'application/json',
        'token': token,
    }

    data = {
        "dbName":"db102",
        "tagName": tag_name,
        "startMsTime":start_time,
        "endMsTime":end_time,
        "resampleMode":1,
        "resamplePriodMs":resamplePriodMs,
    }

    response = requests.post(url, json=data, headers=headers)
    parsed_response = parse_json_response(response.text)

    real_data = {}
    if not parsed_response or 'data' not in parsed_response or not parsed_response['data']:
        return jsonify({"error": "No data found for the given tag name"}), 404
    for i in parsed_response['data'][-data_num:]:
        real_data.update({i['t']: i[tag_name]})
    
    df = pd.DataFrame({'timestamp': list(real_data.keys()), 'value': list(real_data.values())})
    df['timestamp'] = pd.to_datetime(df['timestamp'] + 8 * 3600 * 1000, unit='ms').dt.strftime('%Y-%m-%d %H:%M')
    df.set_index('timestamp', inplace=True)
    df = df.round(2)

    return df.to_dict()['value']

@app.route('/getReason', methods=['GET'])
def get_reason():
    """
    获取耗差计算公式，及各参数数据
    """
    hc = request.args.get('hc')
    time = request.args.get('time')

@app.route('/hcfx/top1', methods=['GET'])
def get_top_hcfx():
    """获取实时耗差分析中数据的最大值
    """
    tag_names = HCFX_TAG_NAMES
    real_data = get_real_value(tag_names)
    real_data = {k: v for k, v in real_data.items() if v is not None}  # 过滤掉值为 None 的数据
    if not real_data:
        return jsonify({"error": "No data found"}), 404
    max_values = {k: v for k, v in real_data.items() if v == max(real_data.values())}
    max_values = {HCFX_TAG_NAMES_REVERSE[k]: v for k, v in max_values.items()}
    return jsonify(max_values)

@app.route('/hcfx/history', methods=['POST'])
def get_his_hcfx():
    """获取历史耗差分析数据

    从请求中获取标签名和查询天数，默认查询最近3天的数据。
    """
    tag_name = request.json.get('tagName')
    days = int(request.json.get('days', 3))  # 默认查询最近3天的数据
    now = int(time.time() * 1000)  # 当前时间戳，单位为毫秒
    start_time = now - days * 24 * 3600 * 1000  # 计算开始时间戳，单位为毫秒
    resamplePriodMs = 5 * 60 * 1000  # 设置重采样周期为5分钟

    url = 'http://168.168.10.82:20001/v1/getHisResampleValue'

    token = get_token()

    headers = {
        'Content-Type': 'application/json',
        'token': token,
    }

    data = {
        "dbName":"db102",
        "tagName": tag_name,
        "startMsTime":start_time,
        "endMsTime":now,
        "resampleMode":1,
        "resamplePriodMs":resamplePriodMs,
    }

    response = requests.post(url, json=data, headers=headers)
    parsed_response = parse_json_response(response.text)

    real_data = {}
    if not parsed_response or 'data' not in parsed_response or not parsed_response['data']:
        return jsonify({"error": "No data found for the given tag name"}), 404
    for i in parsed_response['data']:
        real_data.update({i['t']: i[tag_name]})
    
    df = pd.DataFrame({'timestamp': list(real_data.keys()), 'value': list(real_data.values())})
    df['timestamp'] = pd.to_datetime(df['timestamp'] + 8 * 3600 * 1000, unit='ms').dt.strftime('%Y-%m-%d %H:%M')
    df.set_index('timestamp', inplace=True)
    df = df.round(2)

    return df.to_dict()['value']

@app.route('/getTagNames', methods=['POST'])
def get_tag_names():

    keywords = request.json.get('keywords')
    if not keywords:
        return jsonify({"error": "Input is required"}), 400

    with open("/home/zyp/codes/dify/api/powerplant/tag_names.yaml", "r", encoding="utf-8") as f:
        tag_names = yaml.safe_load(f)
    
    # formated_input = '|'.join([i.strip() for i in re.split('[,，]', keywords)])

    # pattern = re.compile(formated_input)
    # result = {k: v for k, v in tag_names.items() if pattern.search(k)}

    # 支持多个关键词
    keyword_list = [i.strip() for i in re.split('[,，]', keywords)]
    scored_results = []
    for k, v in tag_names.items():
        max_score = 0
        for kw in keyword_list:
            # 计算相似度，取key和value中最大相似度
            score = difflib.SequenceMatcher(None, kw, k).ratio()
            if score > max_score:
                max_score = score
        if max_score > 0.6:  # 阈值可调整
            scored_results.append((k, v, max_score))
            
    # 按相似度降序排序
    scored_results.sort(key=lambda x: x[2], reverse=True)

    # 只返回相似度最高的那一组（可能有多个并列）
    if scored_results:
        top_score = scored_results[0][2]
        top_results = {k: v for k, v, s in scored_results if s == top_score}
        return jsonify(top_results)
    else:
        return jsonify({})

@app.route('/generate/word', methods=['POST'])
def generate_word_document():
    try:
        data = request.json or {}
        content = data.get('content', '')
        filename = data.get('filename', 'document.docx')
        opts = data.get('options', {})

        doc = Document()

        # 页边距设置（厘米）
        section = doc.sections[0]
        margins = opts.get('margins', {})
        section.left_margin = Cm(margins.get('left', 2.5))
        section.right_margin = Cm(margins.get('right', 2.5))
        section.top_margin = Cm(margins.get('top', 2.5))
        section.bottom_margin = Cm(margins.get('bottom', 2.5))

        # 页眉设置
        header_text = opts.get('header')
        if header_text:
            hdr = section.header.paragraphs[0]
            hdr.text = header_text
            hdr.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 页脚设置
        footer_text = opts.get('footer')
        if footer_text:
            ftr = section.footer.paragraphs[0]
            ftr.text = footer_text
            ftr.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 标题设置
        title = opts.get('title', '生成的文档')
        title_p = doc.add_heading(title, 0)
        title_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 标题字体设置
        title_font = opts.get('titleFont', 'SimSun')  # 宋体
        title_size = opts.get('titleSize', 18)
        for run in title_p.runs:
            run.font.name = title_font
            run.font.size = Pt(title_size)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), title_font)

        # 正文段落设置
        body_font = opts.get('font', 'SimSun')
        body_size = opts.get('fontSize', 12)
        line_spacing = opts.get('lineSpacing', 1.5)
        align_map = {
            'LEFT': WD_ALIGN_PARAGRAPH.LEFT,
            'CENTER': WD_ALIGN_PARAGRAPH.CENTER,
            'RIGHT': WD_ALIGN_PARAGRAPH.RIGHT,
            'JUSTIFY': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        align = align_map.get(opts.get('align', 'LEFT').upper(), WD_ALIGN_PARAGRAPH.LEFT)

        # 添加正文内容
        for line in content.split('\n'):
            if line.strip():  # 非空行
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = line_spacing
                p.paragraph_format.space_before = Pt(opts.get('spaceBefore', 0))
                p.paragraph_format.space_after = Pt(opts.get('spaceAfter', 6))
                p.paragraph_format.alignment = align

                run = p.add_run(line)
                run.font.name = body_font
                run.font.size = Pt(body_size)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), body_font)
            else:
                doc.add_paragraph()  # 空行

        # 表格设置（可选）
        table_conf = opts.get('table')
        if table_conf and table_conf.get('data'):
            data_rows = table_conf.get('data', [])
            style = table_conf.get('style', 'Table Grid')
            
            if data_rows:
                table = doc.add_table(rows=len(data_rows), cols=len(data_rows[0]), style=style)
                table.autofit = False
                
                for i, row in enumerate(data_rows):
                    for j, val in enumerate(row):
                        cell = table.cell(i, j)
                        cell.text = str(val)
                        # 设置表格字体
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = body_font
                                run.font.size = Pt(table_conf.get('fontSize', 10))
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), body_font)

        # 分页设置（可选）
        if opts.get('pageBreak'):
            doc.add_page_break()

        # 保存到内存
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return send_file(
            file_stream,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/generate/pdf', methods=['POST'])
def generate_pdf_document():
    try:
        data = request.json or {}
        content = data.get('content', '')
        filename = data.get('filename', 'document.pdf')
        opts = data.get('options', {})

        # 中文字体注册
        font_name = opts.get('font', 'NotoSansCJKsc')
        font_path = opts.get('fontPath', '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc')
        
        # 尝试注册中文字体
        try:
            pdfmetrics.registerFont(TTFont(font_name, font_path))
        except Exception:
            # 字体注册失败时，尝试其他常见字体路径
            font_paths = [
                '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
                '/System/Library/Fonts/PingFang.ttc',
                '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf'
            ]
            for path in font_paths:
                try:
                    pdfmetrics.registerFont(TTFont(font_name, path))
                    break
                except:
                    continue
            else:
                # 所有字体都失败时回退到内置字体
                font_name = 'Helvetica'

        buffer = io.BytesIO()
        
        # 页面设置
        page_size = A4 if opts.get('pageSize', 'A4') == 'A4' else letter
        doc = SimpleDocTemplate(
            buffer,
            pagesize=page_size,
            leftMargin=cm * opts.get('marginLeft', 2),
            rightMargin=cm * opts.get('marginRight', 2),
            topMargin=cm * opts.get('marginTop', 2.5),
            bottomMargin=cm * opts.get('marginBottom', 2.5),
        )

        # 样式设置
        styles = getSampleStyleSheet()
        
        # 标题样式
        title_style = ParagraphStyle(
            'TitleCN', 
            parent=styles['Title'],
            fontName=font_name, 
            fontSize=opts.get('titleSize', 18),
            alignment=TA_CENTER,
            spaceAfter=12
        )
        
        # 正文样式
        align_map = {
            'LEFT': TA_LEFT, 'CENTER': TA_CENTER,
            'RIGHT': TA_RIGHT, 'JUSTIFY': TA_JUSTIFY
        }
        body_style = ParagraphStyle(
            'BodyCN', 
            parent=styles['BodyText'],
            fontName=font_name, 
            fontSize=opts.get('fontSize', 12),
            leading=opts.get('leading', 18),  # 行距
            alignment=align_map.get(opts.get('align', 'LEFT').upper(), TA_LEFT),
            spaceAfter=6
        )

        # 构建文档内容
        story = []
        
        # 添加标题
        title = opts.get('title', '生成的PDF文档')
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 12))

        # 添加正文内容
        for line in content.split('\n'):
            if line.strip():
                story.append(Paragraph(line, body_style))
            else:
                story.append(Spacer(1, 6))  # 空行

        # 添加表格（可选）
        table_conf = opts.get('table')
        if table_conf and table_conf.get('data'):
            data_rows = table_conf.get('data', [])
            if data_rows:
                # 表格样式
                table_style = TableStyle([
                    ('FONTNAME', (0, 0), (-1, -1), font_name),
                    ('FONTSIZE', (0, 0), (-1, -1), table_conf.get('fontSize', 10)),
                    ('ALIGN', (0, 0), (-1, -1), table_conf.get('align', 'LEFT')),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('BACKGROUND', (0, 0), (-1, 0), colors.whitesmoke),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
                ])
                
                t = Table(data_rows, style=table_style)
                story.append(Spacer(1, 12))
                story.append(t)

        # 分页设置（可选）
        if opts.get('pageBreak'):
            story.append(PageBreak())

        # 页眉页脚处理函数
        def on_page(canvas, doc_obj):
            canvas.saveState()
            canvas.setFont(font_name, 9)
            
            # 页眉
            header = opts.get('header')
            if header:
                canvas.drawString(doc_obj.leftMargin, page_size[1] - doc_obj.topMargin + 15, header)
            
            # 页脚
            footer = opts.get('footer', f'第 {doc_obj.page} 页')
            canvas.drawRightString(
                page_size[0] - doc_obj.rightMargin, 
                doc_obj.bottomMargin - 15, 
                footer
            )
            canvas.restoreState()

        # 生成PDF
        doc.build(story, onFirstPage=on_page, onLaterPages=on_page)

        buffer.seek(0)
        return send_file(
            buffer, 
            as_attachment=True, 
            download_name=filename, 
            mimetype='application/pdf'
        )
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    print("=" * 60)
    print("⚠️  DEPRECATED: test.py 已被重构")
    print("请使用新的模块化结构:")
    print("python main.py")
    print("=" * 60)
    app.run(debug=True,host='0.0.0.0',port=6002)
    # app.run(debug=True, host='localhost', port=6002)  # 修改为本地地址和端口
