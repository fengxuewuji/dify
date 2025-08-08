"""
文档生成服务
处理Word和PDF文档生成
"""
import io
from typing import Dict, Any, Optional, List
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.units import cm
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


class DocumentGenerator:
    """文档生成器基类"""
    
    def __init__(self, font_paths: List[str] = None):
        self.font_paths = font_paths or []
    
    def _register_font(self, font_name: str, font_path: str) -> bool:
        """注册字体"""
        try:
            pdfmetrics.registerFont(TTFont(font_name, font_path))
            return True
        except Exception:
            return False


class WordGenerator(DocumentGenerator):
    """Word文档生成器"""
    
    def generate_document(self, content: str, filename: str, options: Dict[str, Any]) -> io.BytesIO:
        """生成Word文档"""
        doc = Document()
        
        # 页边距设置
        section = doc.sections[0]
        margins = options.get('margins', {})
        section.left_margin = Cm(margins.get('left', 2.5))
        section.right_margin = Cm(margins.get('right', 2.5))
        section.top_margin = Cm(margins.get('top', 2.5))
        section.bottom_margin = Cm(margins.get('bottom', 2.5))
        
        # 页眉设置
        header_text = options.get('header')
        if header_text:
            hdr = section.header.paragraphs[0]
            hdr.text = header_text
            hdr.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 页脚设置
        footer_text = options.get('footer')
        if footer_text:
            ftr = section.footer.paragraphs[0]
            ftr.text = footer_text
            ftr.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 标题设置
        title = options.get('title', '生成的文档')
        title_p = doc.add_heading(title, 0)
        title_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 标题字体设置
        title_font = options.get('titleFont', 'SimSun')
        title_size = options.get('titleSize', 18)
        for run in title_p.runs:
            run.font.name = title_font
            run.font.size = Pt(title_size)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), title_font)
        
        # 正文段落设置
        body_font = options.get('font', 'SimSun')
        body_size = options.get('fontSize', 12)
        line_spacing = options.get('lineSpacing', 1.5)
        
        align_map = {
            'LEFT': WD_ALIGN_PARAGRAPH.LEFT,
            'CENTER': WD_ALIGN_PARAGRAPH.CENTER,
            'RIGHT': WD_ALIGN_PARAGRAPH.RIGHT,
            'JUSTIFY': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        align = align_map.get(options.get('align', 'LEFT').upper(), WD_ALIGN_PARAGRAPH.LEFT)
        
        # 添加正文内容 - 支持段落级格式控制
        paragraphs_config = options.get('paragraphs', [])
        
        if paragraphs_config:
            # 使用段落配置模式
            self._add_configured_paragraphs(doc, paragraphs_config, body_font, body_size)
        else:
            # 使用传统模式
            for line in content.split('\\n'):
                if line.strip():
                    p = doc.add_paragraph()
                    p.paragraph_format.line_spacing = line_spacing
                    p.paragraph_format.space_before = Pt(options.get('spaceBefore', 0))
                    p.paragraph_format.space_after = Pt(options.get('spaceAfter', 6))
                    p.paragraph_format.alignment = align
                    
                    run = p.add_run(line)
                    run.font.name = body_font
                    run.font.size = Pt(body_size)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), body_font)
                else:
                    doc.add_paragraph()
        
        # 表格设置
        table_conf = options.get('table')
        if table_conf and table_conf.get('data'):
            data_rows = table_conf.get('data', [])
            style = table_conf.get('style', 'Table Grid')
            
            if data_rows:
                table = doc.add_table(rows=len(data_rows), cols=len(data_rows[0]), style=style)
                table.autofit = False
                
                for i, row in enumerate(data_rows):
                    for j, val in enumerate(row):
                        cell = table.cell(i, j)
                        cell.text = str(val)
                        # 设置表格字体
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = body_font
                                run.font.size = Pt(table_conf.get('fontSize', 10))
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), body_font)
        
        # 分页设置
        if options.get('pageBreak'):
            doc.add_page_break()
        
        # 保存到内存
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return file_stream
    
    def _add_configured_paragraphs(self, doc: Document, paragraphs_config: List[Dict], 
                                 default_font: str, default_size: int):
        """添加配置的段落"""
        from docx.enum.dml import MSO_THEME_COLOR_INDEX
        
        for para_config in paragraphs_config:
            content = para_config.get('content', '')
            if not content.strip():
                doc.add_paragraph()  # 空段落
                continue
                
            # 创建段落
            if para_config.get('style') == 'heading':
                # 标题段落
                level = para_config.get('level', 1)
                p = doc.add_heading(content, level)
            elif para_config.get('numbered'):
                # 编号段落
                p = doc.add_paragraph(content, style='List Number')
            elif para_config.get('bulleted'):
                # 项目符号段落
                p = doc.add_paragraph(content, style='List Bullet')
            else:
                # 普通段落
                p = doc.add_paragraph()
                run = p.add_run(content)
                
                # 字体设置
                font_name = para_config.get('font', default_font)
                font_size = para_config.get('fontSize', default_size)
                run.font.name = font_name
                run.font.size = Pt(font_size)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                
                # 字体样式
                if para_config.get('bold'):
                    run.font.bold = True
                if para_config.get('italic'):
                    run.font.italic = True
                if para_config.get('underline'):
                    run.font.underline = True
                
                # 字体颜色
                color = para_config.get('color')
                if color:
                    if isinstance(color, str) and color.startswith('#'):
                        # 十六进制颜色
                        from docx.shared import RGBColor
                        rgb = tuple(int(color[i:i+2], 16) for i in (1, 3, 5))
                        run.font.color.rgb = RGBColor(*rgb)
            
            # 段落格式设置
            align_map = {
                'LEFT': WD_ALIGN_PARAGRAPH.LEFT,
                'CENTER': WD_ALIGN_PARAGRAPH.CENTER,
                'RIGHT': WD_ALIGN_PARAGRAPH.RIGHT,
                'JUSTIFY': WD_ALIGN_PARAGRAPH.JUSTIFY
            }
            
            if para_config.get('align'):
                p.paragraph_format.alignment = align_map.get(
                    para_config['align'].upper(), WD_ALIGN_PARAGRAPH.LEFT
                )
            
            if para_config.get('lineSpacing'):
                p.paragraph_format.line_spacing = para_config['lineSpacing']
                
            if para_config.get('spaceBefore'):
                p.paragraph_format.space_before = Pt(para_config['spaceBefore'])
                
            if para_config.get('spaceAfter'):
                p.paragraph_format.space_after = Pt(para_config['spaceAfter'])
                
            # 缩进设置
            if para_config.get('leftIndent'):
                p.paragraph_format.left_indent = Cm(para_config['leftIndent'])
            if para_config.get('rightIndent'):
                p.paragraph_format.right_indent = Cm(para_config['rightIndent'])
            if para_config.get('firstLineIndent'):
                p.paragraph_format.first_line_indent = Cm(para_config['firstLineIndent'])


class PDFGenerator(DocumentGenerator):
    """PDF文档生成器"""
    
    def generate_document(self, content: str, filename: str, options: Dict[str, Any]) -> io.BytesIO:
        """生成PDF文档"""
        # 字体注册
        font_name = options.get('font', 'NotoSansCJKsc')
        font_path = options.get('fontPath', '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc')
        
        # 尝试注册中文字体
        if not self._register_font(font_name, font_path):
            # 尝试其他字体路径
            for path in self.font_paths:
                if self._register_font(font_name, path):
                    break
            else:
                # 所有字体都失败时回退到内置字体
                font_name = 'Helvetica'
        
        buffer = io.BytesIO()
        
        # 页面设置
        page_size = A4 if options.get('pageSize', 'A4') == 'A4' else letter
        doc = SimpleDocTemplate(
            buffer,
            pagesize=page_size,
            leftMargin=cm * options.get('marginLeft', 2),
            rightMargin=cm * options.get('marginRight', 2),
            topMargin=cm * options.get('marginTop', 2.5),
            bottomMargin=cm * options.get('marginBottom', 2.5),
        )
        
        # 样式设置
        styles = getSampleStyleSheet()
        
        # 标题样式
        title_style = ParagraphStyle(
            'TitleCN',
            parent=styles['Title'],
            fontName=font_name,
            fontSize=options.get('titleSize', 18),
            alignment=TA_CENTER,
            spaceAfter=12
        )
        
        # 正文样式
        align_map = {
            'LEFT': TA_LEFT, 'CENTER': TA_CENTER,
            'RIGHT': TA_RIGHT, 'JUSTIFY': TA_JUSTIFY
        }
        body_style = ParagraphStyle(
            'BodyCN',
            parent=styles['BodyText'],
            fontName=font_name,
            fontSize=options.get('fontSize', 12),
            leading=options.get('leading', 18),
            alignment=align_map.get(options.get('align', 'LEFT').upper(), TA_LEFT),
            spaceAfter=6
        )
        
        # 构建文档内容
        story = []
        
        # 添加标题
        title = options.get('title', '生成的PDF文档')
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 12))
        
        # 添加正文内容 - 支持段落级格式控制
        paragraphs_config = options.get('paragraphs', [])
        
        if paragraphs_config:
            # 使用段落配置模式
            self._add_configured_pdf_content(story, paragraphs_config, styles, font_name)
        else:
            # 使用传统模式
            for line in content.split('\\n'):
                if line.strip():
                    story.append(Paragraph(line, body_style))
                else:
                    story.append(Spacer(1, 6))
        
        # 添加表格
        table_conf = options.get('table')
        if table_conf and table_conf.get('data'):
            data_rows = table_conf.get('data', [])
            if data_rows:
                table_style = TableStyle([
                    ('FONTNAME', (0, 0), (-1, -1), font_name),
                    ('FONTSIZE', (0, 0), (-1, -1), table_conf.get('fontSize', 10)),
                    ('ALIGN', (0, 0), (-1, -1), table_conf.get('align', 'LEFT')),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('BACKGROUND', (0, 0), (-1, 0), colors.whitesmoke),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
                ])
                
                t = Table(data_rows, style=table_style)
                story.append(Spacer(1, 12))
                story.append(t)
        
        # 分页设置
        if options.get('pageBreak'):
            story.append(PageBreak())
        
        # 页眉页脚处理函数
        def on_page(canvas_obj, doc_obj):
            canvas_obj.saveState()
            canvas_obj.setFont(font_name, 9)
            
            # 页眉
            header = options.get('header')
            if header:
                canvas_obj.drawString(doc_obj.leftMargin, page_size[1] - doc_obj.topMargin + 15, header)
            
            # 页脚
            footer = options.get('footer', f'第 {doc_obj.page} 页')
            canvas_obj.drawRightString(
                page_size[0] - doc_obj.rightMargin,
                doc_obj.bottomMargin - 15,
                footer
            )
            canvas_obj.restoreState()
        
        # 生成PDF
        doc.build(story, onFirstPage=on_page, onLaterPages=on_page)
        
        buffer.seek(0)
        return buffer
    
    def _add_configured_pdf_content(self, story: List, paragraphs_config: List[Dict], 
                                  styles, font_name: str):
        """添加配置的PDF段落"""
        numbered_counter = 1
        
        for para_config in paragraphs_config:
            content = para_config.get('content', '')
            if not content.strip():
                story.append(Spacer(1, 6))
                continue
            
            # 创建段落样式
            style_name = f"Custom_{len(story)}"
            
            # 对齐方式
            align_map = {
                'LEFT': TA_LEFT, 'CENTER': TA_CENTER,
                'RIGHT': TA_RIGHT, 'JUSTIFY': TA_JUSTIFY
            }
            alignment = align_map.get(para_config.get('align', 'LEFT').upper(), TA_LEFT)
            
            # 字体和大小
            para_font = para_config.get('font', font_name)
            font_size = para_config.get('fontSize', 12)
            line_height = para_config.get('leading', font_size * 1.2)
            
            # 创建自定义样式
            if para_config.get('style') == 'heading':
                # 标题样式
                level = para_config.get('level', 1)
                base_style = styles['Heading1'] if level == 1 else styles['Heading2']
                custom_style = ParagraphStyle(
                    style_name,
                    parent=base_style,
                    fontName=para_font,
                    fontSize=font_size + (6 - level * 2),  # 标题字号递减
                    alignment=alignment,
                    spaceBefore=para_config.get('spaceBefore', 12),
                    spaceAfter=para_config.get('spaceAfter', 6),
                    leading=line_height
                )
                story.append(Paragraph(content, custom_style))
                
            elif para_config.get('numbered'):
                # 编号段落
                numbered_content = f"{numbered_counter}. {content}"
                numbered_counter += 1
                
                custom_style = ParagraphStyle(
                    style_name,
                    parent=styles['BodyText'],
                    fontName=para_font,
                    fontSize=font_size,
                    alignment=alignment,
                    spaceBefore=para_config.get('spaceBefore', 3),
                    spaceAfter=para_config.get('spaceAfter', 3),
                    leftIndent=para_config.get('leftIndent', 20),
                    leading=line_height
                )
                story.append(Paragraph(numbered_content, custom_style))
                
            elif para_config.get('bulleted'):
                # 项目符号段落
                bullet_content = f"• {content}"
                
                custom_style = ParagraphStyle(
                    style_name,
                    parent=styles['BodyText'],
                    fontName=para_font,
                    fontSize=font_size,
                    alignment=alignment,
                    spaceBefore=para_config.get('spaceBefore', 3),
                    spaceAfter=para_config.get('spaceAfter', 3),
                    leftIndent=para_config.get('leftIndent', 20),
                    leading=line_height
                )
                story.append(Paragraph(bullet_content, custom_style))
                
            else:
                # 普通段落
                custom_style = ParagraphStyle(
                    style_name,
                    parent=styles['BodyText'],
                    fontName=para_font,
                    fontSize=font_size,
                    alignment=alignment,
                    spaceBefore=para_config.get('spaceBefore', 0),
                    spaceAfter=para_config.get('spaceAfter', 6),
                    leftIndent=para_config.get('leftIndent', 0),
                    rightIndent=para_config.get('rightIndent', 0),
                    firstLineIndent=para_config.get('firstLineIndent', 0),
                    leading=line_height
                )
                
                # 处理富文本格式
                if any(para_config.get(attr) for attr in ['bold', 'italic', 'underline', 'color']):
                    # 构建富文本内容
                    formatted_content = self._format_pdf_text(content, para_config)
                    story.append(Paragraph(formatted_content, custom_style))
                else:
                    story.append(Paragraph(content, custom_style))
    
    def _format_pdf_text(self, content: str, para_config: Dict) -> str:
        """格式化PDF文本（支持粗体、斜体等）"""
        formatted = content
        
        if para_config.get('bold'):
            formatted = f"<b>{formatted}</b>"
        if para_config.get('italic'):
            formatted = f"<i>{formatted}</i>"
        if para_config.get('underline'):
            formatted = f"<u>{formatted}</u>"
        
        color = para_config.get('color')
        if color and isinstance(color, str) and color.startswith('#'):
            formatted = f'<font color="{color}">{formatted}</font>'
            
        return formatted
